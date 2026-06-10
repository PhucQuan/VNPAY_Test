from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "23110146_TRAN_HOANG_PHUC_QUAN_BTVN.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
      tc_w = OxmlElement("w:tcW")
      tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


doc = Document()
setup_styles(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Bài tập về nhà Nhóm 2 - Tích hợp VNPAY")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(11, 37, 69)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("MSSV: 23110146    Họ tên: Trần Hoàng Phúc Quân").italic = True

doc.add_heading("Phần 1. Câu hỏi lý thuyết", level=1)

doc.add_heading("Câu 1. Các tham số bắt buộc khi tạo Payment URL", level=2)
doc.add_paragraph(
    "Khi người dùng bấm nút thanh toán, backend tạo URL sang cổng VNPAY bằng cách nối các tham số giao dịch, sắp xếp và ký bảo mật. Một số tham số quan trọng gồm:"
)
add_table(
    doc,
    ["Tham số", "Ý nghĩa"],
    [
        ("vnp_Version", "Phiên bản API VNPAY đang sử dụng, ví dụ 2.1.0."),
        ("vnp_Command", "Lệnh xử lý giao dịch, thông thường là pay để yêu cầu thanh toán."),
        ("vnp_TmnCode", "Mã website/merchant do VNPAY cấp, dùng để định danh đơn vị bán hàng."),
        ("vnp_Amount", "Số tiền thanh toán đã nhân 100 theo quy ước của VNPAY."),
        ("vnp_CurrCode", "Loại tiền tệ của giao dịch, thường là VND."),
        ("vnp_TxnRef", "Mã đơn hàng duy nhất do hệ thống website tạo để đối soát giao dịch."),
        ("vnp_OrderInfo", "Nội dung thanh toán hiển thị cho người dùng và phục vụ tra cứu."),
        ("vnp_ReturnUrl", "URL để VNPAY chuyển trình duyệt người dùng về sau khi thanh toán."),
        ("vnp_IpAddr", "Địa chỉ IP của khách hàng thực hiện giao dịch."),
        ("vnp_CreateDate", "Thời điểm tạo giao dịch theo định dạng yyyyMMddHHmmss."),
        ("vnp_SecureHash", "Chữ ký bảo mật để VNPAY kiểm tra dữ liệu không bị sửa đổi."),
    ],
    [1900, 7460],
)

doc.add_heading("Câu 2. Cơ chế bảo mật và mã băm vnp_SecureHash", level=2)
doc.add_paragraph(
    "vnp_SecureHash được tạo bằng cách lấy toàn bộ tham số gửi sang VNPAY, loại bỏ chính tham số vnp_SecureHash nếu có, sắp xếp tên tham số theo thứ tự alphabet, chuyển thành chuỗi query string chuẩn, sau đó ký bằng thuật toán HMAC-SHA512 với khóa bí mật vnp_HashSecret do VNPAY cấp."
)
doc.add_paragraph("Công thức khái quát:")
code = doc.add_paragraph()
code.paragraph_format.left_indent = Inches(0.25)
code.add_run("secureHash = HMAC_SHA512(vnp_HashSecret, sortedQueryString)").font.name = "Consolas"
doc.add_paragraph(
    "Bắt buộc phải có chữ ký này vì dữ liệu giao dịch đi qua trình duyệt và mạng Internet nên có thể bị sửa tham số như số tiền, mã đơn hàng hoặc mã merchant. Khi có HMAC, bên nhận có thể tính lại chữ ký từ dữ liệu nhận được. Nếu chữ ký không khớp thì chứng tỏ dữ liệu không còn toàn vẹn hoặc không do bên có HashSecret hợp lệ tạo ra, request phải bị từ chối."
)

doc.add_heading("Câu 3. Phân biệt Return URL và IPN URL", level=2)
add_table(
    doc,
    ["Tiêu chí", "Return URL", "IPN URL / Webhook"],
    [
        ("Cách hoạt động", "VNPAY redirect trình duyệt của khách hàng về website.", "VNPAY gọi trực tiếp server-to-server đến backend của website."),
        ("Mục đích chính", "Hiển thị kết quả thanh toán cho người dùng.", "Thông báo kết quả chính thức để backend cập nhật đơn hàng."),
        ("Phụ thuộc người dùng", "Có. Nếu người dùng tắt trình duyệt hoặc mất mạng, redirect có thể không tới website.", "Không phụ thuộc trình duyệt của người dùng."),
        ("Độ tin cậy cập nhật database", "Không nên dùng làm nguồn cập nhật duy nhất.", "Nên dùng làm cơ chế chính để cập nhật trạng thái thanh toán."),
    ],
    [1700, 3830, 3830],
)
doc.add_paragraph(
    "Trong tình huống khách hàng thanh toán thành công trên điện thoại nhưng mất mạng hoặc tắt trình duyệt ngay sau đó, hệ thống phải dùng IPN URL để đảm bảo database vẫn cập nhật đơn hàng thành 'Đã thanh toán'. Lý do là IPN là request từ server VNPAY sang server website, không phụ thuộc việc trình duyệt khách hàng có quay lại Return URL hay không."
)

doc.add_heading("Câu 4. Validation trước khi UPDATE trạng thái đơn hàng", level=2)
doc.add_paragraph("Trước khi cập nhật database, backend cần kiểm tra theo các bước:")
for item in [
    "Lấy toàn bộ tham số VNPAY gửi về, loại bỏ vnp_SecureHash và vnp_SecureHashType trước khi tính lại chữ ký.",
    "Sắp xếp tham số đúng quy định và tính lại HMAC-SHA512 bằng vnp_HashSecret.",
    "So sánh chữ ký tự tính với vnp_SecureHash nhận được; nếu sai thì từ chối request.",
    "Kiểm tra vnp_TmnCode có đúng mã merchant của hệ thống hay không.",
    "Kiểm tra vnp_TxnRef có tồn tại trong database và đúng đơn hàng đang chờ thanh toán.",
    "Kiểm tra vnp_Amount có đúng số tiền của đơn hàng, lưu ý VNPAY gửi số tiền đã nhân 100.",
    "Kiểm tra mã phản hồi vnp_ResponseCode và vnp_TransactionStatus đều là 00 trước khi coi là thành công.",
    "Kiểm tra đơn hàng chưa được cập nhật thanh toán trước đó để tránh xử lý lặp IPN.",
    "Ghi log các request thất bại hoặc bất thường để phục vụ điều tra và đối soát.",
]:
    add_bullet(doc, item)

doc.add_heading("Phần 2. Bài tập thực hành", level=1)
doc.add_paragraph(
    "Project demo được đặt trong thư mục vnpay-homework-group2. Source code được push lên GitHub public, nhưng không đưa file .env chứa TmnCode và HashSecret thật lên repository."
)
doc.add_heading("Các chức năng đã có trong source demo", level=2)
for item in [
    "Trang chủ có form nhập số tiền và nút Thanh toán qua VNPAY.",
    "Backend tạo Payment URL với các tham số VNPAY và chữ ký HMAC-SHA512.",
    "Trang /return xác thực Secure Hash và hiển thị kết quả giao dịch.",
    "Trang kết quả hiển thị mã đơn hàng, số tiền, ngân hàng, mã giao dịch VNPAY, thời gian thanh toán và mã phản hồi.",
    "Endpoint /ipn đóng vai trò webhook để VNPAY gọi server-to-server và cập nhật trạng thái đơn hàng.",
    "File .env.example hướng dẫn cấu hình Ngrok, ReturnUrl và IPN URL.",
]:
    add_bullet(doc, item)

doc.add_heading("Thông tin nộp bài cần điền", level=2)
add_numbered(doc, "File Word nộp bài: 23110146_TRAN_HOANG_PHUC_QUAN_BTVN.docx.")
add_numbered(doc, "Push thư mục vnpay-homework-group2 lên GitHub public.")
add_numbered(doc, "Link GitHub: https://github.com/PhucQuan/VNPAY_Test")
add_numbered(doc, "Quay video demo khoảng 30 giây gồm: bấm thanh toán, chuyển sang VNPAY sandbox, nhập thẻ test, quay lại website và thấy bảng dữ liệu trả về.")

doc.save(OUTPUT)
print(OUTPUT)
